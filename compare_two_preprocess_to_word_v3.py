from pathlib import Path
import argparse
import random
import tempfile

from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"}


def list_images(root: Path):
    return sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in IMAGE_EXTS)


def build_map(root: Path):
    return {p.relative_to(root).as_posix(): p for p in list_images(root)}


def select_keys(left_map, sources, per_source, random_mode=False, seed=42):
    rng = random.Random(seed)
    selected = []

    for source in sources:
        prefix = source.strip("/\\") + "/"
        keys = sorted(k for k in left_map.keys() if k.startswith(prefix))

        if random_mode:
            rng.shuffle(keys)

        chosen = keys[:per_source]
        selected.extend(chosen)

        print(f"{source}: tìm thấy {len(keys)} ảnh, lấy {len(chosen)} ảnh")

    return selected


def ensure_rgb_white_bg(img: Image.Image) -> Image.Image:
    if img.mode in ("RGBA", "LA"):
        bg = Image.new("RGB", img.size, "white")
        alpha = img.split()[-1]
        bg.paste(img.convert("RGB"), mask=alpha)
        return bg
    return img.convert("RGB")


def prepare_image(src: Path, dst: Path, scale=12, invert=False):
    img = Image.open(src)
    img = ensure_rgb_white_bg(img)

    if invert:
        img = ImageOps.invert(img)

    w, h = img.size
    img = img.resize((max(1, w * scale), max(1, h * scale)), Image.Resampling.NEAREST)

    # Viền để dễ nhìn trong Word
    img = ImageOps.expand(img, border=2, fill="black")
    img = ImageOps.expand(img, border=8, fill="white")

    dst.parent.mkdir(parents=True, exist_ok=True)
    img.save(dst)


def set_cell_text(cell, text, bold=False, size=10, color=(0, 0, 0)):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)


def add_note(cell, text, color=(90, 90, 90)):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*color)


def add_picture(cell, path: Path, width=2.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def create_report(
    left_dir: Path,
    right_dir: Path,
    out_docx: Path,
    sources,
    per_source,
    random_mode,
    seed,
    invert_left=True,
    invert_right=False,
    left_title="Preprocess 1",
    right_title="Preprocess 2",
    scale=12,
    image_width=2.2,
):
    left_map = build_map(left_dir)
    right_map = build_map(right_dir)

    keys = select_keys(left_map, sources, per_source, random_mode=random_mode, seed=seed)

    common = sum(1 for k in keys if k in right_map)
    missing = len(keys) - common

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    title = doc.add_heading("BÁO CÁO SO SÁNH KẾT QUẢ TIỀN XỬ LÝ ẢNH", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Thư mục preprocess 1: ").bold = True
    p.add_run(str(left_dir))
    p.add_run("\nThư mục preprocess 2: ").bold = True
    p.add_run(str(right_dir))
    p.add_run("\nNguồn lấy mẫu: ").bold = True
    p.add_run(", ".join(sources))
    p.add_run("\nSố ảnh mỗi nguồn: ").bold = True
    p.add_run(str(per_source))
    p.add_run("\nTổng ảnh đưa vào Word: ").bold = True
    p.add_run(str(len(keys)))
    p.add_run("\nSố ảnh có đủ cả hai bên: ").bold = True
    p.add_run(str(common))
    p.add_run("\nSố ảnh preprocess2 bị thiếu: ").bold = True
    p.add_run(str(missing))

    note = doc.add_paragraph()
    note.add_run("Ghi chú: ").bold = True
    note.add_run(
        "Cột preprocess1 mặc định được đảo màu khi hiển thị để đưa về nền trắng - nét đen. "
        "Việc đảo màu chỉ áp dụng trong file Word, không sửa ảnh gốc."
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    set_cell_text(header[0], left_title, bold=True, size=11)
    set_cell_text(header[1], right_title, bold=True, size=11)
    shade_cell(header[0], "D9EAF7")
    shade_cell(header[1], "D9EAF7")

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        for idx, key in enumerate(keys, start=1):
            row = table.add_row().cells
            for cell in row:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            left_path = left_map[key]
            right_path = right_map.get(key)

            left_tmp = tmp / f"{idx:05d}_left.png"
            prepare_image(left_path, left_tmp, scale=scale, invert=invert_left)
            add_picture(row[0], left_tmp, width=image_width)
            add_note(row[0], f"{idx}. {key}")
            if invert_left:
                add_note(row[0], "Đã đảo màu để hiển thị")

            if right_path:
                right_tmp = tmp / f"{idx:05d}_right.png"
                prepare_image(right_path, right_tmp, scale=scale, invert=invert_right)
                add_picture(row[1], right_tmp, width=image_width)
                add_note(row[1], f"{idx}. {key}")
            else:
                set_cell_text(
                    row[1],
                    "Không có ảnh tương ứng\ntrong preprocess2",
                    bold=True,
                    size=10,
                    color=(180, 0, 0),
                )
                add_note(row[1], f"{idx}. {key}", color=(180, 0, 0))
                add_note(row[1], "Preprocess2 lỗi hoặc không xuất ảnh", color=(180, 0, 0))
                shade_cell(row[1], "FCE4D6")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print("Đã tạo file Word:", out_docx)


def main():
    parser = argparse.ArgumentParser(
        description="Tạo Word so sánh: lấy 100 ảnh data1 + 100 ảnh data2."
    )
    parser.add_argument("--left", default="output_preprocess1")
    parser.add_argument("--right", default="output_preprocess2")
    parser.add_argument("--out", default="so_sanh_tien_xu_ly_200_mau.docx")
    parser.add_argument("--sources", nargs="+", default=["data1", "data2"])
    parser.add_argument("--per-source", type=int, default=100)
    parser.add_argument("--random", action="store_true", help="Lấy ngẫu nhiên thay vì lấy theo thứ tự")
    parser.add_argument("--seed", type=int, default=42)
    parser.add_argument("--left-title", default="Preprocess 1")
    parser.add_argument("--right-title", default="Preprocess 2")
    parser.add_argument("--no-invert-left", action="store_true")
    parser.add_argument("--invert-right", action="store_true")
    parser.add_argument("--scale", type=int, default=12)
    parser.add_argument("--image-width", type=float, default=2.2)
    args = parser.parse_args()

    left_dir = Path(args.left)
    right_dir = Path(args.right)

    if not left_dir.exists():
        raise FileNotFoundError(f"Không tìm thấy thư mục left: {left_dir}")
    if not right_dir.exists():
        raise FileNotFoundError(f"Không tìm thấy thư mục right: {right_dir}")

    create_report(
        left_dir=left_dir,
        right_dir=right_dir,
        out_docx=Path(args.out),
        sources=args.sources,
        per_source=args.per_source,
        random_mode=args.random,
        seed=args.seed,
        invert_left=not args.no_invert_left,
        invert_right=args.invert_right,
        left_title=args.left_title,
        right_title=args.right_title,
        scale=args.scale,
        image_width=args.image_width,
    )


if __name__ == "__main__":
    main()
